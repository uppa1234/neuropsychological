from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import warnings
warnings.filterwarnings('ignore')
from flask import Flask, request, send_file, render_template
import tempfile


def link(l:list[str]) -> str:
    text = ''
    for i, w in enumerate(l):
        text += w
        if i == len(l) - 1:
            break
        elif i == len(l) - 2:
            text += ' และ '
        else:
            text += ', '
    return text


def get_mmse(df):
    score = int(df.query('Topic == "MMSE" & Subtopic == "Main"').Value.item())
    text = f'Score {score}/30  '

    # Preserved
    if score == 30:
        return text
    
    # Impaired
    text += ' เสียใน '

    impaired = df.query('Topic == "MMSE" & Value != How & Subtopic != "Main"')
    imp_scores = []

    for _, item in impaired.iterrows():
        imp_scores.append(f'{item.Item} {int(item.Value)}/{int(item.How)}')

    imp_test = link(imp_scores)
    text += imp_test

    return text

def get_preserved(list, topic, df):
    preserved = []
    for item in list:
        try:
            print(df.query(f'Topic == "{topic}" & Item == "{item}"'))
            query = df.query(f'Topic == "{topic}" & Item == "{item}"').Interpretation.item() == 'Preserved'
        except:
            raise ValueError(f'Item {item} not found in {topic}')
        if query:
            preserved.append(item)
    if len(preserved) == 0:
        return 'Preserved: None'
    return 'Preserved: ' + link(preserved)

def get_preserved_long(topic, df):
    preserved_items = df.query(f'Topic == "{topic}" & Interpretation == "Preserved"')

    preserved_main = preserved_items.query('Subtopic == "Main"').Item.unique().tolist()
    preserved_sub = preserved_items.query('Subtopic != "Main"').Item.unique().tolist()

    if not preserved_main and not preserved_sub:
        return 'Preserved: None'

    if not preserved_main:
        return f'Preserved: ใน subtest ของ {link(preserved_sub)}'

    preserved_main = [p for p in preserved_main if p != 'Full-Scale IQ']

    return f'Preserved: {link(preserved_main)}\n(ใน subtest ของ {link(preserved_sub)})'


def get_impaired(topic: str, df) -> str:
    imp = df.query(f'Topic == "{topic}" & Subtopic != "Main" & Interpretation != "Preserved"')
    imp_l = imp.Item.unique().tolist()
    imp_m = df.query(f'Topic == "{topic}" & Subtopic == "Main" & Interpretation != "Preserved"')
    imp_m_l = imp_m.Item.unique().tolist()
    if len(imp_l) == 0:
        return 'Impairment: None'
    if len(imp_m_l) == 0:
        return f'Impairment: ใน subtest ของ {link(imp_l)}'
    return f'Impairment: {link(imp_m_l)}\n(ใน subtest ของ {link(imp_l)})'

def get_gca(df):

    # Impairment
    imp_text = get_impaired('General Cognitive Ability', df=df)

    # Preserved
    parts = ['Verbal Parts', 'Performance Parts']
    # pre_text = get_preserved(parts, 'General Cognitive Ability', df=df)
    pre_text = get_preserved_long('General Cognitive Ability', df=df)

    # Full-scale
    fsiq = df.query('Item == "Full-Scale IQ"').Interpretation.item()
    fsiq_text = f'สรุปในภาพรวมพบว่า {fsiq} ใน Full scale IQ'


    return [imp_text, pre_text, fsiq_text]

def get_attn(df):

    # Impairment
    imp_text = get_impaired('Attention', df=df)

    # Preserved
    parts = ['Simple Attention', 'Complex Attention']
    pre_text = get_preserved(parts, 'Attention', df=df)

    return [imp_text, pre_text]

def get_exec(df):

    # Impairment
    imp_text = get_impaired('Executive Function', df=df)

    # Preserved
    parts = ['Stroop', 'Letter Fluency', 'Category Fluency', 'Tower Test', 'Trail Making Test', 'Motor Speed'] # , 'Clock Drawing', 'Number Letter Switching'
    pre_text = get_preserved(parts, 'Executive Function', df=df)

    return [imp_text, pre_text]

def get_lang(df):

    # Impairment
    imp_text = get_impaired('Language', df=df)

    # Preserved
    parts = ['Category Fluency', 'Boston Naming']
    pre_text = get_preserved(parts, 'Language', df=df)

    return [imp_text, pre_text]

def get_vis(df):

    # Impairment
    imp_text = get_impaired('Visuoconstruction', df=df)

    # Preserved
    parts = ['Basic visuoconstruction', 'Complex visuoconstruction']
    pre_text = get_preserved(parts, 'Visuoconstruction', df=df)

    return [imp_text, pre_text]

def get_mem(df):

    # Impairment
    imp_text = get_impaired('Memory', df=df)

    # Preserved
    parts = ['Immediate Memory', 'Delayed Memory', 'General Memory', 'Delayed Auditory Recognition', 'Working Memory', 'Learning Slope', 'Retention', 'Retrieval']
    # pre_text = get_preserved(parts, 'Memory', df=df)
    pre_text = get_preserved_long('Memory', df=df)

    return [imp_text, pre_text]


def generate_word_document(excel_path):

    # output_path = excel_path.with_stem(excel_path.stem + '_summary').with_suffix('.docx')  #TODO adjust path to include name
    # output_path = "/mnt/data/generated_doc.docx"
    temp_dir = tempfile.gettempdir()
    output_path = f"{temp_dir}/generated_doc.docx"
    print(f"Input file path: {excel_path}")
    print(f"Output file path: {output_path}")

    # Read df

    df = pd.read_excel(excel_path, -1)
    TOPICS = ['General Cognitive Ability', 'Attention', 'Memory', 'Executive Function', 'Language', 'Visuoconstruction']
    df.Item = df.Item.str.strip(' ')
    topics = df.Topic.unique().tolist()


    doc = Document()

    # Header

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    df_header = pd.read_excel(excel_path, 0)
    header_text = (
        df_header.iloc[21, 2]
    )
    # header_text = header_text[:header_text.find('ครั้งก่อนทำวันที่')].replace(' ', '  ')
    header_paragraph.text = header_text
    header_run = header_paragraph.runs[0]
    header_run.bold = True
    heading = doc.add_heading('Neuropsychological Test Report', level=1)

    # Center the heading
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Make the heading black
    heading_run = heading.runs[0]
    heading_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color

    doc.add_paragraph()

    # MMSE
    mmse = doc.add_paragraph()
    mmse_run = mmse.add_run('MMSE: ')
    mmse_run.bold = True
    mmse.add_run(get_mmse(df=df))

    # General Cognitive Ability
    gca = doc.add_paragraph()
    gca_run = gca.add_run('General Cognitive ability (WAIS-III):')
    gca_run.bold = True
    for bullet in get_gca(df=df):
        doc.add_paragraph(bullet, style='ListBullet')

    # Attention
    attn = doc.add_paragraph()
    attn_run = attn.add_run('Attention:')
    attn_run.bold = True
    for bullet in get_attn(df=df):
        doc.add_paragraph(bullet, style='ListBullet')

    # Memory
    mem = doc.add_paragraph()
    mem_run = mem.add_run('Memory:')
    mem_run.bold = True
    for bullet in get_mem(df=df):
        doc.add_paragraph(bullet, style='ListBullet')

    # Executive Function
    exec = doc.add_paragraph()
    exec_run = exec.add_run('Executive Function:')
    exec_run.bold = True
    for bullet in get_exec(df=df):
        doc.add_paragraph(bullet, style='ListBullet')

    # Language
    lang = doc.add_paragraph()
    lang_run = lang.add_run('Language:')
    lang_run.bold = True
    for bullet in get_lang(df=df):
        doc.add_paragraph(bullet, style='ListBullet')

    # Visuoconstruction
    vis = doc.add_paragraph()
    vis_run = vis.add_run('Visuoconstruction:')
    vis_run.bold = True
    for bullet in get_vis(df=df):
        doc.add_paragraph(bullet, style='ListBullet')
        
    # Summary

    # Logic
    sig_imp = df.query('Interpretation == "Significant impairment" & Topic != "MMSE"')
    just_imp = df.query('Interpretation != "Preserved" & Interpretation != "Significant impairment" & Topic != "MMSE"')
    sig_imp_l = sig_imp['Topic'].unique().tolist() # start with these first but not used separately
    just_imp_l = just_imp['Topic'].unique().tolist()
    just_imp_l = [j for j in just_imp_l if j not in sig_imp_l]
    preserved = [p for p in TOPICS if p not in sig_imp_l + just_imp_l]

    summary = doc.add_paragraph()
    summary_run = summary.add_run('Summary: ')
    summary_run.bold = True
    doc.add_paragraph('Impairment: ' + ' ในบางส่วนของ ' + link(sig_imp_l + just_imp_l), style='ListBullet')
    doc.add_paragraph('Preserved: ' + link(preserved), style='ListBullet')
    doc.add_paragraph('Profile นี้ _______________ ทั้งนี้ผลต้องใช้ clinical information ประกอบร่วมด้วย', style='ListBullet')


    # Save the document
    doc.save(output_path)
    
    print("Word document has been created successfully.")

    return output_path


# App

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return "No file uploaded", 400

    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400

    # Process the Excel file
    try:
        word_file_path = generate_word_document(file)
        return send_file(word_file_path, as_attachment=True)
    except Exception as e:
        return str(e), 500


if __name__ == '__main__':
    app.run(debug=True)
